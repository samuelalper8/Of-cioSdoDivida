import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import zipfile
import os
from datetime import datetime, timedelta
import unicodedata
import re

# --- CONFIGURA√á√ÉO DA P√ÅGINA ---
st.set_page_config(page_title="Gerador de Of√≠cios - ConPrev", layout="wide")

# ================= 1. FUN√á√ïES DE LIMPEZA E NORMALIZA√á√ÉO =================

def remove_accents(input_str):
    if not isinstance(input_str, str): return str(input_str)
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])

def normalize_key_nospace(text):
    if pd.isna(text): return ""
    text = remove_accents(str(text)).upper().strip()
    prefixes = ["MUNICIPIO DE ", "PREFEITURA DE ", "PREFEITURA MUNICIPAL DE ", "CAMARA MUNICIPAL DE ", "FUNDO MUNICIPAL DE "]
    for p in prefixes:
        if text.startswith(p): text = text[len(p):]
    return text.replace(" ", "").replace("_", "").replace("-", "")

def normalize_key_standard(text):
    if pd.isna(text): return ""
    text = remove_accents(str(text)).upper().strip()
    prefixes = ["MUNICIPIO DE ", "PREFEITURA DE ", "PREFEITURA MUNICIPAL DE ", "FUNDO MUNICIPAL DE "]
    for p in prefixes:
        if text.startswith(p): text = text[len(p):].strip()
    return text

def format_camel_case(text):
    if not text: return ""
    if " " in text: return text
    return re.sub(r'(?<!^)(?=[A-Z])', ' ', text)

def extrair_uf_filename(nome_arquivo):
    if not isinstance(nome_arquivo, str): return "GO"
    parts = nome_arquivo.replace(" ", "_").split('_')
    ufs = {'AC','AL','AP','AM','BA','CE','DF','ES','GO','MA','MT','MS','MG','PA','PB','PR','PE','PI','RJ','RN','RS','RO','RR','SC','SP','SE','TO'}
    for part in parts:
        if part.upper() in ufs:
            return part.upper()
    parts_dash = nome_arquivo.split('-')
    if len(parts_dash) > 0 and parts_dash[0].strip().upper() in ufs:
        return parts_dash[0].strip().upper()
    return "GO"

# ================= 2. CARREGAMENTO DE DADOS =================

def gerar_modelo_responsaveis():
    data = {'Munic√≠pio': ['Goi√¢nia', 'An√°polis'], 'Respons√°vel': ['Prefeito A', 'Prefeito B']}
    return pd.DataFrame(data).to_csv(index=False, sep=';', encoding='utf-8-sig').encode('utf-8-sig')

def gerar_modelo_pgfn():
    data = {'Arquivo': ['GO_Cidade_PGFN.pdf'], 'Identificador': ['123'], 'Modalidade': ['Lei 13.485'], 'Saldo (R$)': ['1000.00']}
    return pd.DataFrame(data).to_csv(index=False, sep=',', encoding='utf-8-sig').encode('utf-8-sig')

def carregar_responsaveis(arquivo):
    """
    Carrega a lista de respons√°veis priorizando a coluna 'Nome Extra√≠do'.
    """
    try:
        if arquivo.name.endswith('.csv'):
            try: df = pd.read_csv(arquivo, sep=';', encoding='utf-8-sig')
            except: 
                arquivo.seek(0)
                df = pd.read_csv(arquivo, sep=',', encoding='utf-8')
        else:
            df = pd.read_excel(arquivo)
        
        # Normaliza colunas
        df.columns = [remove_accents(c).strip().lower() for c in df.columns]
        
        # --- L√ìGICA DE PRIORIDADE DE COLUNAS ---
        
        # 1. Tenta achar especificamente "Nome Extraido" (Sua planilha V2)
        col_resp = next((c for c in df.columns if 'nome extraido' in c), None)
        
        # 2. Se n√£o achar, tenta "Respons√°vel", "Prefeito" (Planilha V1)
        if not col_resp:
            col_resp = next((c for c in df.columns if any(k in c for k in ['responsavel', 'prefeito', 'gestor'])), None)
            
        # 3. Para cidade/√≥rg√£o
        col_muni = next((c for c in df.columns if any(k in c for k in ['orgao', 'entidade', 'municipio', 'cidade'])), None)
        
        if not col_muni or not col_resp:
            st.error(f"‚ö†Ô∏è N√£o foi poss√≠vel identificar as colunas na lista de respons√°veis. Colunas lidas: {list(df.columns)}")
            return {}
        
        dic = {}
        for _, row in df.iterrows():
            val_muni = str(row[col_muni])
            val_resp = str(row[col_resp]).strip()
            
            # Cria chaves normalizadas para facilitar o encontro
            # Ex: "MUNICIPIO DE ALMAS" -> "ALMAS"
            dic[normalize_key_standard(val_muni)] = val_resp
            dic[normalize_key_nospace(val_muni)] = val_resp
            
        return dic
    except Exception as e:
        st.error(f"Erro ao processar respons√°veis: {e}")
        return {}

def carregar_pgfn_csv(arquivo):
    try:
        if arquivo.name.endswith('.csv'): df = pd.read_csv(arquivo)
        else: df = pd.read_excel(arquivo)
        
        cols = {c.lower(): c for c in df.columns}
        col_arq = cols.get('arquivo', df.columns[0])
        col_id = next((c for c in df.columns if 'identificador' in c.lower() or 'processo' in c.lower()), None)
        col_mod = next((c for c in df.columns if 'modalidade' in c.lower()), None)
        col_val = next((c for c in df.columns if 'saldo' in c.lower() or 'valor' in c.lower()), None)

        dados = {}
        meta = {}

        for _, row in df.iterrows():
            nome_arq = str(row[col_arq])
            parts = nome_arq.split('_')
            cidade_raw = "DESCONHECIDO"
            uf_raw = "GO"
            
            if len(parts) >= 2:
                if len(parts[0]) == 2:
                    uf_raw = parts[0].upper()
                    cidade_raw = parts[1]
                else:
                    cidade_raw = parts[0]
            
            key = normalize_key_nospace(cidade_raw)
            if key not in meta: meta[key] = {'UF': uf_raw, 'Nome': cidade_raw}
            if key not in dados: dados[key] = []
            
            val_str = str(row[col_val]).replace('R$', '').replace('.', '').replace(',', '.')
            try: val_float = float(val_str)
            except: val_float = 0.0

            dados[key].append({
                'Processo': str(row[col_id]),
                'Modalidade': str(row[col_mod]),
                'Valor Original': val_float,
                'Fonte': 'PGFN CSV'
            })
        return dados, meta
    except Exception as e:
        st.error(f"Erro ao ler PGFN: {e}")
        return {}, {}

def buscar_responsavel(muni_display, key_nospace, db_resp):
    # 1. Busca exata pela chave sem espa√ßos
    if key_nospace in db_resp: return db_resp[key_nospace]
    
    # 2. Busca pelo nome normalizado
    norm_std = normalize_key_standard(muni_display)
    if norm_std in db_resp: return db_resp[norm_std]
    
    # 3. Busca aproximada
    for k in db_resp:
        if k.startswith(norm_std) or norm_std.startswith(k):
            return db_resp[k]
            
    return "PREFEITO(A) MUNICIPAL"

# ================= 3. MANIPULA√á√ÉO WORD =================

def replace_everywhere(doc: Document, old: str, new: str) -> None:
    def repl(par):
        if old in par.text:
            for run in par.runs:
                if old in run.text: run.text = run.text.replace(old, new)
            if old in par.text: par.text = par.text.replace(old, new)
    for p in doc.paragraphs: repl(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: repl(p)
    for s in doc.sections:
        for h in [s.header, s.first_page_header, s.footer, s.first_page_footer]:
            if h:
                for p in h.paragraphs: repl(p)

def formatar_valor(val):
    if isinstance(val, (int, float)):
        return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return str(val)

def adicionar_linha_tabela(table, orgao, modalidade, processo, valor, is_placeholder=False):
    row_cells = table.add_row().cells
    
    p1 = row_cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.add_run(orgao)
    
    if is_placeholder:
        run_nc = p1.add_run("\n(nada consta)")
        run_nc.font.size = Pt(8)
    elif modalidade and modalidade.lower() != 'nan':
        p1.add_run(f"\n({modalidade})").font.size = Pt(8)

    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if processo and processo.lower() != 'nan':
        p2.add_run(str(processo))
    else:
        p2.add_run("-")

    p3 = row_cells[2].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(valor)

    for cell in row_cells:
        cell.vertical_alignment = 1
        for p in cell.paragraphs:
            if p.runs and p.runs[0].font.size != Pt(8):
                p.runs[0].font.size = Pt(10)

def preencher_tabela(table, df_rfb, lista_pgfn):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False 

    hdr_cells = table.rows[0].cells
    titulos = ['√ìrg√£o / Modalidade', 'Processo / Documento', 'Saldo em 31/12/2025']
    for i, t in enumerate(titulos):
        hdr_cells[i].text = t
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            if not p.runs: p.add_run(t).font.bold = True

    # RFB
    if not df_rfb.empty:
        df_clean = df_rfb[~df_rfb['Sistema'].astype(str).str.contains("PGFN", case=False, na=False)]
        df_validas = df_clean.dropna(subset=['Valor Original'])
        
        has_valid_row = False
        if not df_validas.empty:
            for _, row in df_validas.iterrows():
                if row['Valor Original'] == 0 and pd.isna(row['Processo']): continue
                has_valid_row = True
                mod = str(row['Modalidade']) if pd.notna(row['Modalidade']) else ""
                proc = str(row['Processo']) if pd.notna(row['Processo']) else "-"
                adicionar_linha_tabela(table, "Receita Federal do Brasil", mod, proc, formatar_valor(row['Valor Original']))
        
        if not has_valid_row:
            adicionar_linha_tabela(table, "Receita Federal do Brasil", "", "-", "-", is_placeholder=True)
    else:
        adicionar_linha_tabela(table, "Receita Federal do Brasil", "", "-", "-", is_placeholder=True)

    # PGFN
    if lista_pgfn:
        for item in lista_pgfn:
            adicionar_linha_tabela(table, "Procuradoria Geral da Fazenda Nacional", item['Modalidade'], item['Processo'], formatar_valor(item['Valor Original']))
    else:
        df_pgfn_ex = df_rfb[df_rfb['Sistema'].astype(str).str.contains("PGFN", case=False, na=False)] if not df_rfb.empty else pd.DataFrame()
        if not df_pgfn_ex.empty:
             for _, row in df_pgfn_ex.iterrows():
                mod = str(row['Modalidade']) if pd.notna(row['Modalidade']) else ""
                proc = str(row['Processo']) if pd.notna(row['Processo']) else "-"
                adicionar_linha_tabela(table, "Procuradoria Geral da Fazenda Nacional", mod, proc, formatar_valor(row['Valor Original']))
        else:
            adicionar_linha_tabela(table, "Procuradoria Geral da Fazenda Nacional", "", "-", "-", is_placeholder=True)

def inserir_tabela_no_placeholder(doc, df_rfb, lista_pgfn, placeholder="{{TABELA}}"):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = ""
            table = doc.add_table(rows=1, cols=3)
            paragraph._p.addnext(table._tbl)
            preencher_tabela(table, df_rfb, lista_pgfn)
            return True
    return False

# ================= 4. INTERFACE =================
st.title("Gerador de Of√≠cios 8.1 (Corre√ß√£o Respons√°vel)")

with st.expander("üìÇ Baixar Modelos"):
    c1, c2 = st.columns(2)
    c1.download_button("üì• Modelo Respons√°veis", gerar_modelo_responsaveis(), "Modelo_Responsaveis.csv", "text/csv")
    c2.download_button("üì• Modelo Extrato PGFN", gerar_modelo_pgfn(), "Modelo_Extrato_PGFN.csv", "text/csv")

col1, col2 = st.columns(2)
with col1:
    uploaded_excel = st.file_uploader("1. D√≠vidas RFB (Excel)", type=["xlsx"])
    uploaded_pgfn = st.file_uploader("4. Extrato PGFN (CSV/Excel)", type=["csv", "xlsx"])
with col2:
    uploaded_template = st.file_uploader("2. Modelo Word (.docx)", type=["docx"])
    uploaded_resp = st.file_uploader("3. Lista Respons√°veis (CSV)", type=["csv"])

st.sidebar.header("Par√¢metros")
num_inicial = st.sidebar.number_input("N¬∫ Inicial", value=46)
ano_doc = st.sidebar.number_input("Ano", value=2026)

# ================= 5. PROCESSAMENTO =================
if st.button("üöÄ Gerar Arquivos"):
    if not uploaded_template:
        st.error("Modelo Word √© obrigat√≥rio.")
        st.stop()
    if not uploaded_excel and not uploaded_pgfn:
        st.error("Envie pelo menos uma fonte de dados.")
        st.stop()
    
    db_resp = carregar_responsaveis(uploaded_resp) if uploaded_resp else {}
    dados_pgfn, meta_pgfn = carregar_pgfn_csv(uploaded_pgfn) if uploaded_pgfn else ({}, {})
    
    if uploaded_resp and not db_resp:
        st.warning("‚ö†Ô∏è Planilha de respons√°veis lida, mas nenhuma informa√ß√£o foi extra√≠da. Verifique os t√≠tulos das colunas.")

    try:
        df_rfb = pd.DataFrame()
        if uploaded_excel:
            df_rfb = pd.read_excel(uploaded_excel, engine='openpyxl')
            col_muni = 'Munic√≠pio' if 'Munic√≠pio' in df_rfb.columns else df_rfb.columns[0]
            col_arq = 'Arquivo' if 'Arquivo' in df_rfb.columns else None
            df_rfb = df_rfb.dropna(subset=[col_muni])
            df_rfb[col_muni] = df_rfb[col_muni].astype(str).str.strip()
            df_rfb['Key'] = df_rfb[col_muni].apply(normalize_key_nospace)
            
            def get_uf_rfb(row):
                if col_arq and pd.notna(row[col_arq]): return extrair_uf_filename(row[col_arq])
                return "GO"
            df_rfb['UF_Ref'] = df_rfb.apply(get_uf_rfb, axis=1)

        mestre_munis = {}
        if not df_rfb.empty:
            for _, row in df_rfb.iterrows():
                k = row['Key']
                if k not in mestre_munis:
                    mestre_munis[k] = {'Nome': row[col_muni], 'UF': row['UF_Ref']}
        
        for k, meta in meta_pgfn.items():
            if k not in mestre_munis:
                nome_formatado = format_camel_case(meta['Nome'])
                mestre_munis[k] = {'Nome': nome_formatado, 'UF': meta['UF']}

        munis_por_uf = {}
        for k, dados in mestre_munis.items():
            uf = dados['UF']
            if uf not in munis_por_uf: munis_por_uf[uf] = []
            munis_por_uf[uf].append({'Key': k, 'Nome': dados['Nome']})
        
        ufs_ordenadas = sorted(munis_por_uf.keys())
        contador = num_inicial
        logs = []
        
        hoje = datetime.now() - timedelta(hours=3)
        meses = {1:"janeiro", 2:"fevereiro", 3:"mar√ßo", 4:"abril", 5:"maio", 6:"junho",
                 7:"julho", 8:"agosto", 9:"setembro", 10:"outubro", 11:"novembro", 12:"dezembro"}
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for uf in ufs_ordenadas:
                lista_alvo = sorted(munis_por_uf[uf], key=lambda x: x['Nome'])
                for item in lista_alvo:
                    key = item['Key']
                    nome_display = item['Nome']
                    
                    df_rfb_muni = df_rfb[df_rfb['Key'] == key] if not df_rfb.empty else pd.DataFrame()
                    lista_pgfn_muni = dados_pgfn.get(key, [])
                    
                    nome_pref = buscar_responsavel(nome_display, key, db_resp)
                    if nome_pref == "PREFEITO(A) MUNICIPAL" and db_resp:
                        logs.append(f"‚ö†Ô∏è {nome_display} ({uf}): Respons√°vel n√£o encontrado.")

                    uploaded_template.seek(0)
                    doc = Document(uploaded_template)
                    
                    replaces = {
                        "{{MUNICIPIO}}": f"{nome_display.upper()} ‚Äì {uf}", 
                        "{{UF}}": uf,
                        "{{PREFEITO}}": nome_pref.upper(),
                        "{{NUM_OFICIO}}": f"{contador:03d}/{ano_doc}",
                        "{{DATA_EXTENSO}}": f"Goi√¢nia, {hoje.day} de {meses[hoje.month]} de {hoje.year}."
                    }
                    for k_rep, v_rep in replaces.items(): replace_everywhere(doc, k_rep, v_rep)
                    
                    sucesso = inserir_tabela_no_placeholder(doc, df_rfb_muni, lista_pgfn_muni, "{{TABELA}}")
                    if not sucesso: inserir_tabela_no_placeholder(doc, df_rfb_muni, lista_pgfn_muni, "{{TABELA_DEBITOS}}")

                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    fname = f"{uf}/{contador:03d}-{ano_doc} - {uf} - {nome_display} - Saldo Divida RFB-PGFN.docx"
                    zf.writestr(fname, doc_io.getvalue())
                    
                    contador += 1
            
        st.success(f"‚úÖ Processo Finalizado! {contador - num_inicial} of√≠cios gerados.")
        if logs:
            with st.expander("Alertas de Processamento"):
                for l in logs: st.write(l)

        st.download_button(
            label="‚¨áÔ∏è Baixar TODOS os Of√≠cios (Pacote Completo)",
            data=zip_buffer.getvalue(),
            file_name=f"Oficios_Completos_{hoje.strftime('%Y%m%d')}.zip",
            mime="application/zip"
        )

    except Exception as e:
        st.error(f"Erro Cr√≠tico: {e}")
